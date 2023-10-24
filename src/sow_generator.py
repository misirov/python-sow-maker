from __future__ import print_function
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
import os
from docx import Document
from datetime import datetime


class SOWGenerator:
    """
    Generates a Statement Of Work (SOW) document by fetching values from a Google Spreadsheet.

    Requirements:
        - Google Spreadsheet API access.
        - 'credentials.json' file for authentication.
          Upon first authentication, a 'token.json' will be generated for subsequent API calls.

    Methods:
        - cantina_managed(SPREADSHEET_ID): Generates a Cantina Managed SOW.
        - guild(SPREADSHEET_ID): Generates a GUILD SOW.
    """

    def cantina_managed(SPREADSHEET_ID):
        """
        Generates a Cantina Managed SOW document by fetching specific values from a Google Spreadsheet.

        Args:
            SPREADSHEET_ID (str): ID of the Google Spreadsheet to fetch data from.
        """
        print("Creating Cantina Managed SOW...")
        RANGE_NAME = 'cm!A1:E27'  # Define the range of cells to retrieve data from.

        spreadsheet = Spreadsheet(SPREADSHEET_ID, RANGE_NAME)  # Initialize the Spreadsheet class.
        data = spreadsheet.fetch_data()  # Retrieve the spreadsheet data.

        # Convert the data into a dictionary suitable for replacing placeholders in the docx template.
        replacements = {}
        for row in data:
            for i in range(len(row)-1):
                key, value = row[i], row[i+1]
                if key and value:
                    replacements[key] = value
        formatted_replacements = {"{" + key + "}": value for key, value in replacements.items()}

        print("\nCreated dictionary to replace for placeholders\n\n", formatted_replacements)

        # Process the SOW template and save the filled document.
        template_path = "templates/sow_template.docx"
        sow_filler = SOWFiller(template_path)
        sow_filler.replace_placeholders(formatted_replacements)


    def guild(SPREADSHEET_ID):
        """
        Generates a GUILD SOW document by fetching specific values from a Google Spreadsheet.

        Args:
            SPREADSHEET_ID (str): ID of the Google Spreadsheet to fetch data from.
        """
        print("Creating GUILD SOW...")
        RANGE_NAME = 'guild!A1:E17'  # Define the range of cells to retrieve data from.

        spreadsheet = Spreadsheet(SPREADSHEET_ID, RANGE_NAME)  # Initialize the Spreadsheet class.
        data = spreadsheet.fetch_data()  # Retrieve the spreadsheet data.

        # Convert the data into a dictionary suitable for replacing placeholders in the docx template.
        replacements = {}
        for row in data:
            for i in range(len(row)-1):
                key, value = row[i], row[i+1]
                if key and value:
                    replacements[key] = value
        formatted_replacements = {"{" + key + "}": value for key, value in replacements.items()}

        print("\nCreated dictionary to replace for placeholders\n\n", formatted_replacements)

        # Process the SOW template and save the filled document.
        template_path = "templates/guild_template.docx"
        sow_filler = SOWFiller(template_path)
        sow_filler.replace_placeholders(formatted_replacements)



class SOWFiller:
    """ The SOWFiller class processes DOCX templates to replace specified placeholders with given values. 
    It supports modifications in the main text, tables, headers, and footers, 
    then saves the updated document in an "output" directory. 
    """

    def __init__(self, template_path):
        """Constructor with path to the template document"""
        self.template_path = template_path


    def replace_in_paragraphs(self, paragraphs, replacements):
        """Takes a paragraph and replaces for KEYWORD:VALUE in replacements dictionary."""
        for paragraph in paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace(key, value)


    def replace_placeholders(self, replacements):
        """Replace all KEYWORD instances in a document and saves modified copy in /output"""
        try:
            doc = Document(self.template_path)

            # Replace in main document body
            self.replace_in_paragraphs(doc.paragraphs, replacements)

            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.replace_in_paragraphs(cell.paragraphs, replacements)

            # Replace in headers and footers
            for section in doc.sections:
                header = section.header
                footer = section.footer
                self.replace_in_paragraphs(header.paragraphs, replacements)
                self.replace_in_paragraphs(footer.paragraphs, replacements)

            # Save modified document
            if not os.path.exists("output"):
                os.makedirs("output")
            
            # Construct the new filename
            base_name = os.path.basename(self.template_path).split('.')[0]
            timestamp = datetime.now().strftime('%Y-%m-%d')
            new_filename = f"{base_name}_{timestamp}.docx"
            print(f"\nSaving modified document as {new_filename} in output/")
            doc.save(f"output/{new_filename}")

        except Exception as e:
            print(f"An error has occurred\n{e}")



class Spreadsheet:
    """ 
    Authenticate and fetch data from the provided spreadsheet.
    
    Attributes:
        SCOPES (list): OAuth 2.0 scope for read-only access to user's sheets.
        spreadsheet_id (str): ID of the target Google Spreadsheet.
        range_name (str): String denoting the range of cells to retrieve.
        service: Authenticated service used to make Google Sheets API calls.
    """

    SCOPES = ['https://www.googleapis.com/auth/spreadsheets.readonly']

    def __init__(self, spreadsheet_id, range_name):
        """
        Initializes an instance of the Spreadsheet class.
        
        Args:
            spreadsheet_id (str): ID of the target Google Spreadsheet.
            range_name (str): String denoting the range of cells to retrieve.
        """
        self.spreadsheet_id = spreadsheet_id
        self.range_name = range_name
        self.service = self.authenticate()


    def authenticate(self):
        """
        Authenticates the user with Google Sheets using OAuth 2.0.
        
        Checks for an existing 'token.json', otherwise prompts the user
        for authorization and saves the token.

        Returns:
            service: Authenticated service to make Google Sheets API calls.
        """
        creds = None
        if os.path.exists('token.json'):
            creds = Credentials.from_authorized_user_file('token.json', self.SCOPES)

        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file('credentials.json', self.SCOPES)
                creds = flow.run_local_server(port=0)
            with open('token.json', 'w') as token:
                token.write(creds.to_json())
        return build('sheets', 'v4', credentials=creds)


    def fetch_data(self):
        """
        Fetches and prints the data from the specified Google Spreadsheet range.

        Returns:
            list: List of lists representing the spreadsheet rows and cells.
            Returns an empty list if no data is found or an error occurs.
        """
        print("fetching data from spreadsheet...")
        try:
            sheet = self.service.spreadsheets()
            result = sheet.values().get(spreadsheetId=self.spreadsheet_id, range=self.range_name).execute()
            values = result.get('values', [])
            print(f"\nPrinting retrieved values \n\n{values}")
            if not values:
                print('No data found.')
                return []
            return values
        except HttpError as err:
            print(err)
            return []
