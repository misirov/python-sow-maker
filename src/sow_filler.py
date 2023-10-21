import os
from docx import Document

class SOWFiller:
    def __init__(self, template_path):
        self.template_path = template_path


    def replace_in_paragraphs(self, paragraphs, replacements):
        for paragraph in paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace(key, value)

    # Replacements
    def replace_placeholders(self, replacements):
        try:
            doc = Document(self.template_path)

            # Replace in main document body
            self.replace_in_paragraphs(doc.paragraphs, replacements)

            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.replace_in_paragraphs(cell.paragraphs, replacements)

            # Replace in headers and footers
            for section in doc.sections:
                header = section.header
                footer = section.footer
                self.replace_in_paragraphs(header.paragraphs, replacements)
                self.replace_in_paragraphs(footer.paragraphs, replacements)

            # Save modified document
            if not os.path.exists("output"):
                os.makedirs("output")
            print("saving modified template in output/ ...")
            doc.save("output/modified.docx")

        except Exception as e:
            print(f"An error has occurred\n{e}")
