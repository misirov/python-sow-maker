from docx import Document
from datetime import datetime, timedelta

def replace_in_paragraphs(paragraphs, replacements):
    for paragraph in paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

def replace_placeholders(doc_path, replacements):
    doc = Document(doc_path)

    # Replace in main document body
    replace_in_paragraphs(doc.paragraphs, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, replacements)

    # Replace in headers and footers
    for section in doc.sections:
        header = section.header
        footer = section.footer
        replace_in_paragraphs(header.paragraphs, replacements)
        replace_in_paragraphs(footer.paragraphs, replacements)

    # Save modified document
    print("saving document...")
    doc.save("output/modified.docx")

if __name__ == "__main__":
    today = datetime.today()
    day_issued = today.strftime("%m-%d-%Y")
    five_days_later = today + timedelta(days=5)
    day_expires = five_days_later.strftime("%m-%d-%Y")

    replacements = {
        "{DAY_ISSUED}"                  : day_issued,
        "{DAY_EXPIRES}"                 : day_expires,
        "{CLIENT}"                      : "SUPERCLIENT",
        "{TOTAL_FEE}"                   : "69,420",
        "{MARKETPLACE_FEE}"             : "20",
        "{PAYMENT_DUE_DATE}"            : "TODAY",
        "{SECURITY_REVIEW_START_DATE}"  : "NOT TODAY",
        "{REVIEW_PERIOD_WEEKS}"         : "2.5",
        "{FIX_PERIOD}"                  : "N/A",
        "{SCOPE}"                       : "https://github.com/organization/superrepo-finance/src/Contract.sol",
        "{LSR_NAMES}"                   : "Riley Holterhus, Rajeev",
        "{SR_NAMES}"                    : "M4rio.eth, Sterim",
        "{ASR_NAMES}"                   : "0x4non, Shung",
        "{JSR_NAMES}"                   : "0xChristos, Sabnock",
        "{LSR_AMOUNT}"                  : "2",
        "{LSR_RATE}"                    : "20,000",
        "{LSR_WEEKS}"                   : "5",
        "{LSR_FEE}"                     : "100,000"







    }

    doc_path = "templates/sow_template.docx"
    replace_placeholders(doc_path, replacements)
